"""Generate TAMS-branded DOCX reports using python-docx."""

import os
import re
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

from app.config import get_settings

settings = get_settings()
TAMS_LOGO = os.path.join(settings.ASSETS_DIR, "tams_logo.png")
TAMS_FOOTER = os.path.join(settings.ASSETS_DIR, "tams_footer.png")

# Brand color constants (from TAM Capital Brand Guidelines)
DARK_BLUE_RGB = RGBColor(0x22, 0x2F, 0x62)    # Deep Blue - primary
ACCENT_RGB = RGBColor(0x1A, 0x6D, 0xB6)       # Light Blue - accent
TURQUOISE_RGB = RGBColor(0x6C, 0xB9, 0xB6)    # Turquoise - secondary
GRAY_RGB = RGBColor(0x4A, 0x4A, 0x4A)         # Body text
WHITE_RGB = RGBColor(0xFF, 0xFF, 0xFF)
SOFT_CARBON_RGB = RGBColor(0xB1, 0xB3, 0xB6)  # Neutral
LIGHT_BLUE_HEX = "E8F0F8"                      # Light background for rows
DARK_BLUE_HEX = "222F62"                        # Primary for headers
# Legacy alias
GREEN_RGB = ACCENT_RGB


def _set_cell_shading(cell, color_hex):
    """Set cell background color."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}" w:val="clear"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def _add_header_logo(doc):
    """Add TAMS logo to the document header."""
    section = doc.sections[0]
    header = section.header
    header.is_linked_to_previous = False

    if os.path.exists(TAMS_LOGO):
        paragraph = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = paragraph.add_run()
        run.add_picture(TAMS_LOGO, width=Inches(1.8))


def _add_footer(doc):
    """Add TAMS footer with company details."""
    section = doc.sections[0]
    footer = section.footer
    footer.is_linked_to_previous = False

    if os.path.exists(TAMS_FOOTER):
        paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run()
        run.add_picture(TAMS_FOOTER, width=Inches(5.5))


def _add_cover_page(doc, stock_name, ticker, date_str):
    """Add a styled cover page."""
    # Company name
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.space_after = Pt(0)
    run = p.add_run(stock_name.upper())
    run.font.size = Pt(28)
    run.font.color.rgb = DARK_BLUE_RGB
    run.bold = True
    run.font.name = "Arial"

    # Ticker
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.space_after = Pt(24)
    run = p.add_run(f"({ticker})")
    run.font.size = Pt(16)
    run.font.color.rgb = GREEN_RGB
    run.font.name = "Arial"

    # Report title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.space_after = Pt(6)
    run = p.add_run("Comprehensive Investor Report")
    run.font.size = Pt(18)
    run.font.color.rgb = DARK_BLUE_RGB
    run.font.name = "Arial"

    # Subtitle
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.space_after = Pt(36)
    run = p.add_run("Equity Analysis, News Impact Assessment\n& Geopolitical Risk Assessment")
    run.font.size = Pt(13)
    run.font.color.rgb = GRAY_RGB
    run.font.name = "Arial"

    # Metadata table
    meta_data = [
        ("Date:", date_str),
        ("Prepared by:", "TAM Capital"),
        ("Classification:", "Confidential - For Investor Use Only"),
    ]

    for label, value in meta_data:
        p = doc.add_paragraph()
        p.space_after = Pt(2)
        run = p.add_run(label + "  ")
        run.bold = True
        run.font.size = Pt(11)
        run.font.color.rgb = DARK_BLUE_RGB
        run.font.name = "Arial"
        run = p.add_run(value)
        run.font.size = Pt(11)
        run.font.color.rgb = GRAY_RGB
        run.font.name = "Arial"

    # Page break
    doc.add_page_break()


def _parse_markdown_table(text):
    """Parse a markdown table into rows of cells."""
    lines = text.strip().split("\n")
    rows = []
    for line in lines:
        line = line.strip()
        if not line or line.startswith("|") and set(line.replace("|", "").strip()) <= {"-", ":"}:
            continue
        if line.startswith("|"):
            cells = [c.strip() for c in line.split("|")[1:-1]]
            rows.append(cells)
    return rows


def _add_table_from_markdown(doc, table_text):
    """Create a styled table from markdown table text."""
    rows = _parse_markdown_table(table_text)
    if not rows or len(rows) < 2:
        return

    num_cols = len(rows[0])
    table = doc.add_table(rows=len(rows), cols=num_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    for i, row_data in enumerate(rows):
        row = table.rows[i]
        for j, cell_text in enumerate(row_data):
            if j < num_cols:
                cell = row.cells[j]
                cell.text = ""
                p = cell.paragraphs[0]
                run = p.add_run(cell_text)
                run.font.size = Pt(9)
                run.font.name = "Arial"

                if i == 0:  # Header row
                    _set_cell_shading(cell, DARK_BLUE_HEX)
                    run.font.color.rgb = WHITE_RGB
                    run.bold = True
                else:
                    run.font.color.rgb = GRAY_RGB
                    if i % 2 == 0:
                        _set_cell_shading(cell, LIGHT_BLUE_HEX)

    doc.add_paragraph()  # Space after table


def _add_section_content(doc, content, charts=None):
    """Parse markdown-like content and add to document with proper formatting."""
    lines = content.split("\n")
    in_table = False
    table_lines = []
    i = 0

    while i < len(lines):
        line = lines[i]

        # Check for table start
        if line.strip().startswith("|") and not in_table:
            in_table = True
            table_lines = [line]
            i += 1
            continue

        # Continue table
        if in_table and line.strip().startswith("|"):
            table_lines.append(line)
            i += 1
            continue

        # End of table
        if in_table and not line.strip().startswith("|"):
            in_table = False
            _add_table_from_markdown(doc, "\n".join(table_lines))
            table_lines = []
            # Don't increment i, process current line

        # Headings — use proper styles for outline levels and PDF bookmarks
        if line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2)
            i += 1
            continue

        if line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=3)
            i += 1
            continue

        if line.startswith("#### "):
            p = doc.add_paragraph()
            p.space_before = Pt(8)
            p.space_after = Pt(4)
            run = p.add_run(line[5:].strip())
            run.font.size = Pt(11)
            run.font.color.rgb = DARK_BLUE_RGB
            run.bold = True
            run.font.name = "Arial"
            i += 1
            continue

        # Bold text lines (e.g., **Tailwinds:**)
        if line.strip().startswith("**") and line.strip().endswith("**"):
            p = doc.add_paragraph()
            run = p.add_run(line.strip().replace("**", ""))
            run.bold = True
            run.font.size = Pt(11)
            run.font.color.rgb = DARK_BLUE_RGB
            run.font.name = "Arial"
            i += 1
            continue

        # Bullet points
        if line.strip().startswith("- ") or line.strip().startswith("* "):
            text = line.strip()[2:]
            p = doc.add_paragraph(style="List Bullet")
            # Handle bold within bullets
            parts = re.split(r'\*\*(.*?)\*\*', text)
            for j, part in enumerate(parts):
                run = p.add_run(part)
                run.font.size = Pt(10)
                run.font.name = "Arial"
                run.font.color.rgb = GRAY_RGB
                if j % 2 == 1:
                    run.bold = True
            i += 1
            continue

        # Numbered items
        match = re.match(r'^(\d+)\.\s+(.*)', line.strip())
        if match:
            text = match.group(2)
            p = doc.add_paragraph(style="List Number")
            parts = re.split(r'\*\*(.*?)\*\*', text)
            for j, part in enumerate(parts):
                run = p.add_run(part)
                run.font.size = Pt(10)
                run.font.name = "Arial"
                run.font.color.rgb = GRAY_RGB
                if j % 2 == 1:
                    run.bold = True
            i += 1
            continue

        # Regular paragraph
        text = line.strip()
        if text:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(6)
            # Handle inline bold
            parts = re.split(r'\*\*(.*?)\*\*', text)
            for j, part in enumerate(parts):
                if not part:
                    continue
                run = p.add_run(part)
                run.font.size = Pt(10)
                run.font.name = "Arial"
                run.font.color.rgb = GRAY_RGB
                if j % 2 == 1:
                    run.bold = True
                    run.font.color.rgb = DARK_BLUE_RGB

        i += 1

    # Handle trailing table
    if in_table and table_lines:
        _add_table_from_markdown(doc, "\n".join(table_lines))


def generate_docx_report(stock_name: str, ticker: str, sections: dict,
                         charts: dict = None, output_dir: str = "output",
                         sources=None) -> str:
    """Generate a complete TAMS-branded DOCX report.

    Args:
        stock_name: Company name
        ticker: Stock ticker symbol
        sections: Dict of section_name -> content text
        charts: Dict of chart_name -> chart file path
        output_dir: Directory to save the output file
        sources: SourceCollector instance for references section

    Returns:
        Path to the generated DOCX file
    """
    os.makedirs(output_dir, exist_ok=True)

    doc = Document()

    # Set default font
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(10)
    style.font.color.rgb = GRAY_RGB
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.15

    # Configure heading styles for consistent formatting
    for level, (size, spacing_before, spacing_after) in {
        1: (18, 24, 12),
        2: (16, 18, 10),
        3: (13, 12, 6),
    }.items():
        heading_style = doc.styles[f"Heading {level}"]
        heading_style.font.name = "Arial"
        heading_style.font.size = Pt(size)
        heading_style.font.color.rgb = DARK_BLUE_RGB
        heading_style.font.bold = True
        heading_style.paragraph_format.space_before = Pt(spacing_before)
        heading_style.paragraph_format.space_after = Pt(spacing_after)
        heading_style.paragraph_format.keep_with_next = True

    # Configure list styles
    for list_style_name in ("List Bullet", "List Number"):
        if list_style_name in doc.styles:
            ls = doc.styles[list_style_name]
            ls.font.name = "Arial"
            ls.font.size = Pt(10)
            ls.font.color.rgb = GRAY_RGB
            ls.paragraph_format.space_after = Pt(3)
            ls.paragraph_format.line_spacing = 1.15

    # Page margins
    section = doc.sections[0]
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

    # Add header and footer
    _add_header_logo(doc)
    _add_footer(doc)

    # Cover page
    date_str = datetime.now().strftime("%B %d, %Y")
    _add_cover_page(doc, stock_name, ticker, date_str)

    # Section titles mapping
    section_titles = {
        "executive_summary": "Executive Summary",
        "fundamental_analysis": "Part I: Fundamental Analysis",
        "dividend_analysis": "Dividend Income Analysis",
        "earnings_analysis": "Earnings Analysis",
        "risk_assessment": "Risk Assessment Framework",
        "technical_analysis": "Technical Analysis Dashboard",
        "sector_rotation": "Sector Rotation Strategy",
        "news_impact": "Part II: Recent News Impact Assessment",
        "war_impact": "Part III: Geopolitical Risk Assessment",
        "key_takeaways": "Key Takeaways & Investment Thesis",
    }

    # Add each section
    section_order = [
        "executive_summary", "fundamental_analysis", "dividend_analysis",
        "earnings_analysis", "risk_assessment", "technical_analysis",
        "sector_rotation", "news_impact", "war_impact", "key_takeaways"
    ]

    for section_key in section_order:
        if section_key not in sections:
            continue

        content = sections[section_key]

        # Section header — use Heading 1 for proper outline levels
        title = section_titles.get(section_key, section_key.replace("_", " ").title())
        doc.add_heading(title.upper(), level=1)

        # Section content
        _add_section_content(doc, content)

        # Add relevant chart after certain sections
        if charts:
            chart_map = {
                "fundamental_analysis": ["revenue_earnings", "valuation"],
                "technical_analysis": ["technical"],
                "dividend_analysis": ["dividend"],
            }
            for chart_key in chart_map.get(section_key, []):
                chart_path = charts.get(chart_key, "")
                if chart_path and os.path.exists(chart_path):
                    doc.add_paragraph()  # Space
                    p = doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = p.add_run()
                    run.add_picture(chart_path, width=Inches(5.5))
                    doc.add_paragraph()  # Space

        doc.add_page_break()

    # Sources & References
    if sources:
        doc.add_heading("SOURCES & REFERENCES", level=1)

        if hasattr(sources, 'format_for_docx'):
            source_list = sources.format_for_docx()
            # Group by type
            groups = {}
            for s in source_list:
                stype = {
                    "yahoo_finance": "Market Data (Yahoo Finance)",
                    "financial_statements": "Financial Statements (Yahoo Finance)",
                    "news_article": "News & Analysis",
                    "web_search": "Web Research",
                    "sector_news": "Sector Research",
                }.get(s["type"], "Other Sources")
                if stype not in groups:
                    groups[stype] = []
                groups[stype].append(s)

            for group_name, group_sources in groups.items():
                # Group header
                p = doc.add_paragraph()
                p.space_before = Pt(8)
                run = p.add_run(group_name)
                run.font.size = Pt(11)
                run.font.color.rgb = DARK_BLUE_RGB
                run.bold = True
                run.font.name = "Arial"

                # Sources in group
                for s in group_sources:
                    p = doc.add_paragraph(style="List Bullet")
                    run = p.add_run(f"[{s['index']}] {s['title']}")
                    run.font.size = Pt(9)
                    run.font.color.rgb = GRAY_RGB
                    run.font.name = "Arial"
                    if s.get("url"):
                        p2 = doc.add_paragraph()
                        p2.paragraph_format.left_indent = Cm(1.5)
                        run2 = p2.add_run(s["url"])
                        run2.font.size = Pt(8)
                        run2.font.color.rgb = ACCENT_RGB
                        run2.font.name = "Arial"

                    if s.get("accessed"):
                        p3 = doc.add_paragraph()
                        p3.paragraph_format.left_indent = Cm(1.5)
                        run3 = p3.add_run(f"Accessed: {s['accessed']}")
                        run3.font.size = Pt(8)
                        run3.font.color.rgb = SOFT_CARBON_RGB
                        run3.font.name = "Arial"

        doc.add_page_break()

    # Disclaimer
    doc.add_heading("DISCLAIMER", level=1)

    from app.analysis.compiler import DISCLAIMER_TEXT
    for para_text in DISCLAIMER_TEXT.strip().split("\n\n"):
        p = doc.add_paragraph()
        run = p.add_run(para_text.strip())
        run.font.size = Pt(9)
        run.font.color.rgb = GRAY_RGB
        run.font.name = "Arial"

    # Save
    safe_name = re.sub(r'[^\w\s-]', '', stock_name).strip().replace(" ", "_")
    filename = f"{safe_name}_Investor_Report_TAM_{datetime.now().strftime('%Y%m%d')}.docx"
    filepath = os.path.join(output_dir, filename)
    doc.save(filepath)

    return filepath
